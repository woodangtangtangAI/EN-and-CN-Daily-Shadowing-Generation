import os
import json
import random
import datetime
import asyncio
import tempfile
import edge_tts
from docx import Document
from google import genai
from google.genai import types

# ==========================================
# ⚙️ 설정
# ==========================================
# GitHub Actions에서는 환경변수로 전달, 로컬에서는 직접 입력
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

# GitHub Actions 환경 여부 감지
IS_GITHUB_ACTIONS = os.environ.get("GITHUB_ACTIONS") == "true"

# 1. 언어별 설정
LANGUAGES = [
    {
        "name": "영어",
        "folder": "영어",
        # 듀엣 대화 효과를 위한 여성, 남성 성우진 지정 (생동감 있고 또렷한 톤으로 교체)
        "voices": ["en-US-AriaNeural", "en-US-ChristopherNeural"], 
        "rate": "+0%",                
        "prompt_extra": (
            "미국식 영어 기준으로 번역 및 작성하며, Chunking(끊어 읽기 표기)은 절대 생략하고 자연스러운 원문과 해석만 제공하세요. "
            "제공된 '상황(배경지식)'을 완벽하게 파악하여, 그 상황에 100% 어울리는 실전 TPO(시간, 장소, 상황) 맞춤형 표현을 사용하세요. "
            "뉴스라면 앵커가 쓰는 정제된 표현, 친구와의 대화라면 편안한 슬랭(Slang), 비즈니스 회의라면 격식 있는 프로페셔널한 어휘 등 상황에 따라 문체와 어조를 완벽하게 바꾸세요. "
            "기계적인 직역은 철저히 배제하고, 언어 실력의 뼈대(backbone)가 될 수 있는 네이티브스러운 찐 표현을 지향합니다."
        )
    },
    {
        "name": "중국어",
        "folder": "중국어",
        # 듀엣 대화 효과를 위한 여성, 남성 성우진 지정
        "voices": ["zh-CN-XiaoxiaoNeural", "zh-CN-YunxiNeural"], 
        "rate": "-10%",                  
        "prompt_extra": (
            "중국어 간체자로 번역 및 작성하며, 모든 중국어 원문과 단어에는 반드시 발음 기호인 병음(Pinyin)을 함께 표기하세요. (Chunking 생략) "
            "제공된 '상황(배경지식)'을 완벽하게 파악하여, 그 상황에 100% 어울리는 실전 TPO(시간, 장소, 상황) 맞춤형 표현을 사용하세요. "
            "뉴스라면 앵커톤(播音腔)의 정제된 표현, 친구와의 대화/위챗 메신저라면 리얼 구어체(口语化)와 유행어(예: 吐槽, 搞定), 비즈니스 회의라면 전문 비즈니스 용어 등 상황에 따라 어조를 완벽하게 바꾸세요. "
            "한국식 직역은 절대 피하고, 어설픈 외국인 티를 벗어날 수 있는 네이티브스러운 찐 표현으로 언어 실력의 뼈대(backbone)를 다지게 해주세요."
        )
    }
]

# ==========================================
# 🎲 무한 랜덤 시나리오 풀 (물리적 강제 추출용)
# ==========================================
# AI 자율 선택 시 IT/게임 등 특정 주제로 편향되는 문제를 방지하기 위해,
# Python random.choice()로 매일 1개를 물리적으로 강제 추출하여 프롬프트에 주입합니다.
# 각 항목은 (대분류, 구체적 상황 힌트) 형태입니다.

SCENARIO_POOL = [
    # ===== 일상 대화 / 친구 수다 (구어체 핵심) =====
    ("친구와 카페 수다", "카페에서 친구에게 최근 본 넷플릭스/유튜브 콘텐츠 리뷰를 열정적으로 떠들며 서로 의견이 갈리는 일상 수다"),
    ("소개팅 후기", "소개팅이나 블라인드 데이트를 다녀온 직후 절친에게 솔직한 후기를 털어놓으며 조언을 구하는 대화"),
    ("룸메이트 갈등", "룸메이트나 동거인이 설거지/청소/소음 문제로 서로 서운한 감정을 표현하며 티격태격하는 대화"),
    ("다이어트 고민", "운동이나 다이어트를 시작했다가 작심삼일로 포기한 친구끼리 서로 위로하며 핑계 대는 공감 수다"),
    ("반려동물 일상", "반려견/반려묘의 귀여운 행동이나 병원 방문 에피소드를 친구에게 자랑하거나 걱정을 나누는 대화"),
    ("추억 회상", "오랜만에 만난 친구와 학창시절이나 군대/대학 시절 추억을 떠올리며 그때 그 시절을 회상하는 대화"),
    ("이사/집 구하기", "새 집을 구하러 다니면서 겪은 황당한 에피소드나 부동산 중개인과의 경험을 친구에게 푸는 수다"),
    ("SNS/인스타 논쟁", "친구의 SNS 게시물이나 인플루언서의 논란에 대해 카톡/위챗으로 의견을 나누는 메신저 대화"),

    # ===== 실전 TPO 생활 영어/중국어 =====
    ("카페/레스토랑 주문 실수", "카페나 레스토랑에서 주문이 잘못 나오거나 음식에 문제가 있어 점원에게 정중하게 컴플레인 거는 상황"),
    ("택시/우버 탑승", "택시나 우버를 타고 목적지를 설명하거나 기사와 길이 막혀 짜증 나는 상황에서 대화하는 장면"),
    ("공항/출입국", "공항 체크인, 수하물 분실, 비행기 지연, 게이트 변경 등 여행 중 겪는 예상치 못한 상황 대처"),
    ("호텔 체크인 트러블", "호텔 체크인 시 예약이 안 되어 있거나 방이 기대와 다를 때 프론트 데스크와 나누는 대화"),
    ("병원/약국 방문", "감기나 알레르기로 병원/약국을 방문하여 증상을 설명하고 약 처방을 받는 실전 대화"),
    ("중고거래 네고", "당근마켓이나 온라인 중고거래에서 가격 흥정을 하거나 직거래 장소를 정하는 메시지 대화"),
    ("택배/배달 사고", "택배가 파손되었거나 배달 음식이 오배송/누락되어 고객센터에 전화하는 컴플레인 상황"),

    # ===== 직장/오피스 리얼 대화 =====
    ("탕비실 험담", "회사 탕비실이나 점심시간에 동료끼리 상사의 이해할 수 없는 지시나 야근에 대해 투덜거리는 수다"),
    ("면접 준비/후기", "취업 면접을 앞두고 친구에게 긴장감을 토로하거나, 면접 끝나고 분위기와 질문을 공유하는 대화"),
    ("퇴사 고민", "번아웃이 와서 이직이나 퇴사를 진지하게 고민하며 믿을 수 있는 친구에게 조언을 구하는 대화"),
    ("프레젠테이션 직전", "중요한 발표를 5분 앞두고 동료에게 긴장감을 토로하며 마지막 리허설을 하는 긴박한 대화"),

    # ===== 가족/연애/감정 =====
    ("부모님과 세대 갈등", "부모님과 진로/결혼/생활 방식 등에 대해 의견이 달라 서로 설득하려는 세대 간 대화"),
    ("연인 사이 오해", "연인 사이에서 사소한 오해로 서운한 감정을 표현하거나 화해를 시도하는 감정적 대화"),
    ("형제자매 말다툼", "형제자매끼리 집안일 분담이나 부모님 용돈 등 현실적인 문제로 옥신각신하는 대화"),

    # ===== 취미/문화/엔터테인먼트 =====
    ("영화 리뷰 논쟁", "최근 개봉한 블록버스터 영화나 드라마 시즌 피날레를 보고 친구와 스포일러 섞어가며 열띤 토론"),
    ("콘서트/페스티벌 후기", "콘서트나 뮤직 페스티벌에 다녀온 후 현장 분위기와 감동을 흥분하며 전하는 대화"),
    ("요리 도전기", "유튜브 보고 새로운 요리에 도전했다가 실패한 경험을 친구에게 웃기게 풀어내는 수다"),
    ("운동/헬스장", "헬스장이나 필라테스를 다니기 시작하면서 겪는 근육통, 트레이너와의 에피소드를 나누는 대화"),
    ("게임 협동 플레이", "온라인 게임에서 팀원끼리 작전을 짜거나 실수한 팀원에게 장난스럽게 따지는 인게임 보이스챗"),

    # ===== 사회/시사/뉴스 (구어체로) =====
    ("뉴스 반응 수다", "오늘 본 충격적인 뉴스(경제, 사회, 연예)에 대해 친구끼리 카톡으로 빠르게 의견을 나누는 대화"),
    ("환경/지속가능성", "일회용품 줄이기나 채식, 제로웨이스트 등 환경 이슈에 대해 친구와 가볍게 토론하는 대화"),
    ("AI/신기술 체험기", "ChatGPT, VR, 자율주행 등 신기술을 직접 써보고 신기했던 경험을 친구에게 설명하는 대화"),

    # ===== 여행/해외 생활 =====
    ("해외 여행 해프닝", "해외 여행 중 길을 잃거나 현지인과 소통이 안 되어 겪은 웃기고 황당한 에피소드를 나누는 대화"),
    ("유학/워홀 생활", "해외 유학이나 워킹홀리데이 중 겪는 문화 차이, 향수병, 현지 친구 사귀기 에피소드"),
    ("해외 맛집 탐방", "해외 여행지에서 현지 맛집을 찾아다니며 음식을 평가하고 추천하는 미식 수다"),

    # ===== 쇼핑/소비/재테크 =====
    ("쇼핑 중독 고백", "세일 시즌에 충동구매를 해버린 후 후회하며 친구에게 고백하는 대화"),
    ("재테크/투자 수다", "주식, 부동산, 코인 등 재테크 경험을 친구끼리 가볍게 공유하며 조언을 나누는 대화"),
]

# ==========================================
# 📂 Google Drive 업로드 기능 (GitHub Actions 전용)
# ==========================================
def get_drive_service():
    """Google Drive API 서비스 객체 생성 (OAuth 2.0 사용자 인증)"""
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    
    creds_json = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON") or os.environ.get("GOOGLE_USER_CREDENTIALS")
    if not creds_json:
        print("⚠️ GOOGLE_USER_CREDENTIALS 환경변수가 없습니다. 업로드를 건너뜁니다.")
        return None
    
    creds_info = json.loads(creds_json)
    credentials = Credentials.from_authorized_user_info(
        creds_info, scopes=["https://www.googleapis.com/auth/drive"]
    )
    
    # 만약 토큰이 만료되었다면 갱신
    if credentials.expired and credentials.refresh_token:
        try:
            credentials.refresh(Request())
        except Exception as e:
            print(f"⚠️ 토큰 갱신 중 오류 발생: {e}")
            
    return build("drive", "v3", credentials=credentials)

def find_or_create_folder(service, folder_name, parent_id=None):
    """Google Drive에서 폴더를 찾거나, 없으면 새로 생성"""
    query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
    if parent_id:
        query += f" and '{parent_id}' in parents"
    
    results = service.files().list(q=query, spaces="drive", fields="files(id, name)").execute()
    files = results.get("files", [])
    
    if files:
        return files[0]["id"]
    
    # 폴더가 없으면 새로 생성
    file_metadata = {
        "name": folder_name,
        "mimeType": "application/vnd.google-apps.folder",
    }
    if parent_id:
        file_metadata["parents"] = [parent_id]
    
    folder = service.files().create(body=file_metadata, fields="id").execute()
    print(f"📁 Google Drive 폴더 생성: {folder_name}")
    return folder.get("id")

def upload_to_drive(service, local_path, folder_id, filename):
    """파일을 Google Drive 특정 폴더에 업로드"""
    from googleapiclient.http import MediaFileUpload
    
    # 같은 이름의 파일이 이미 있으면 업데이트
    query = f"name='{filename}' and '{folder_id}' in parents and trashed=false"
    results = service.files().list(q=query, spaces="drive", fields="files(id)").execute()
    existing = results.get("files", [])
    
    media = MediaFileUpload(local_path, resumable=True)
    
    if existing:
        # 기존 파일 업데이트
        service.files().update(fileId=existing[0]["id"], media_body=media).execute()
    else:
        # 새 파일 생성
        file_metadata = {"name": filename, "parents": [folder_id]}
        service.files().create(body=file_metadata, media_body=media, fields="id").execute()
    
    print(f"   ☁️ Google Drive 업로드 완료: {filename}")

def upload_folder_to_drive(service, local_folder, today_folder_name):
    """로컬 출력 폴더의 모든 파일을 Google Drive에 업로드"""
    # [언어 공부] 폴더 명시적 ID (사용자 제공)
    parent_id = "1YmABh3RfKVsqrAVWFelKn2NCtY9ezrCS"
    
    # 오늘 날짜 폴더 찾기/생성
    date_folder_id = find_or_create_folder(service, today_folder_name, parent_id)
    
    # 폴더 안의 모든 파일 업로드
    for filename in os.listdir(local_folder):
        filepath = os.path.join(local_folder, filename)
        if os.path.isfile(filepath):
            upload_to_drive(service, filepath, date_folder_id, filename)


# ==========================================
# ==========================================
# 📄 문서 생성 (Premium Styling System)
# ==========================================
def create_docx(data, filename, lang_name):
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Page Margins (top/bottom: 2.54cm, left/right: 1.91cm)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.91)
        section.right_margin = Cm(1.91)
        
    # Font choice (Microsoft YaHei for Chinese, Malgun Gothic for others)
    base_font = 'Microsoft YaHei' if lang_name == '중국어' else 'Malgun Gothic'
    
    # Document Style
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = base_font
    font_normal.size = Pt(10)
    font_normal.color.rgb = RGBColor(38, 38, 38) # Off-black
    
    # 1. Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(18)
    
    run_title = p_title.add_run(f'🧠 오늘의 {lang_name} 훈련')
    run_title.bold = True
    run_title.font.name = base_font
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(31, 78, 121) # Steel Blue
    
    # 2. Helper functions for sections
    def add_section_heading(emoji, text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        
        run_emoji = h.add_run(f"{emoji} ")
        run_emoji.font.name = base_font
        run_emoji.font.size = Pt(12.5)
        
        run_text = h.add_run(text)
        run_text.bold = True
        run_text.font.name = base_font
        run_text.font.size = Pt(12.5)
        run_text.font.color.rgb = RGBColor(31, 78, 121) # Steel Blue
        
    # --- BACKGROUND SECTION ---
    if data.get('background_kr'):
        add_section_heading("📖", "배경지식")
        p_bg = doc.add_paragraph()
        p_bg.paragraph_format.line_spacing = 1.3
        p_bg.paragraph_format.space_after = Pt(14)
        run_bg = p_bg.add_run(data['background_kr'])
        run_bg.font.name = base_font
        run_bg.font.size = Pt(9.5)
        run_bg.font.color.rgb = RGBColor(90, 90, 90) # Muted Gray
        
    # --- SENTENCES SECTION ---
    if data.get('sentences'):
        add_section_heading("🎧", "청취 훈련 문장")
        
        for i, s in enumerate(data['sentences'], 1):
            # Original Sentence Paragraph
            p_orig = doc.add_paragraph()
            p_orig.paragraph_format.space_before = Pt(6)
            p_orig.paragraph_format.space_after = Pt(0)
            p_orig.paragraph_format.line_spacing = 1.1
            
            run_num = p_orig.add_run(f"{i}. ")
            run_num.bold = True
            run_num.font.name = base_font
            run_num.font.size = Pt(10.5)
            run_num.font.color.rgb = RGBColor(31, 78, 121) # Steel Blue
            
            run_orig = p_orig.add_run(s.get('original', ''))
            run_orig.bold = True
            run_orig.font.name = base_font
            run_orig.font.size = Pt(10.5)
            run_orig.font.color.rgb = RGBColor(0, 0, 0) # Black
            
            # Pinyin Paragraph (Chinese only)
            if s.get('pinyin'):
                p_pinyin = doc.add_paragraph()
                p_pinyin.paragraph_format.left_indent = Inches(0.25)
                p_pinyin.paragraph_format.space_after = Pt(0)
                p_pinyin.paragraph_format.line_spacing = 1.1
                
                run_pinyin = p_pinyin.add_run(f"[{s.get('pinyin')}]")
                run_pinyin.font.name = base_font
                run_pinyin.font.size = Pt(9)
                run_pinyin.font.color.rgb = RGBColor(0, 128, 128) # Teal
                
            # Translation Paragraph
            p_trans = doc.add_paragraph()
            p_trans.paragraph_format.left_indent = Inches(0.25)
            p_trans.paragraph_format.space_after = Pt(10) # Space between sentences
            p_trans.paragraph_format.line_spacing = 1.1
            
            run_trans = p_trans.add_run(f"해석: {s.get('translation', '')}")
            run_trans.font.name = base_font
            run_trans.font.size = Pt(9.5)
            run_trans.font.color.rgb = RGBColor(80, 80, 80) # Muted text
            
    # --- VOCABULARY SECTION ---
    if data.get('vocabulary'):
        add_section_heading("📝", "핵심 단어")
        
        for i, v in enumerate(data['vocabulary'], 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.2
            
            # Number and Word
            run_word = p.add_run(f"{i}. {v.get('word', '')}")
            run_word.bold = True
            run_word.font.name = base_font
            run_word.font.size = Pt(10)
            run_word.font.color.rgb = RGBColor(44, 94, 138) # Medium steel blue
            
            # Pinyin
            if v.get('pinyin'):
                run_py = p.add_run(f" [{v.get('pinyin')}]")
                run_py.font.name = base_font
                run_py.font.size = Pt(9)
                run_py.font.color.rgb = RGBColor(0, 128, 128) # Teal
                
            # Separator & POS & Meaning
            sep_text = " - "
            if v.get('pos'):
                sep_text += f"{v.get('pos')} "
            run_mean = p.add_run(f"{sep_text}{v.get('meaning', '')}")
            run_mean.font.name = base_font
            run_mean.font.size = Pt(9.5)
            run_mean.font.color.rgb = RGBColor(64, 64, 64) # Dark gray

    # --- COLLOCATIONS SECTION ---
    if data.get('collocations'):
        add_section_heading("🧩", "핵심 표현 (Collocations/Idioms)")
        
        for i, c in enumerate(data['collocations'], 1):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.2
            
            # Number and Expression
            run_exp = p.add_run(f"{i}. {c.get('expression', '')}")
            run_exp.bold = True
            run_exp.font.name = base_font
            run_exp.font.size = Pt(10)
            run_exp.font.color.rgb = RGBColor(44, 94, 138)
            
            # Pinyin
            if c.get('pinyin'):
                run_py = p.add_run(f" [{c.get('pinyin')}]")
                run_py.font.name = base_font
                run_py.font.size = Pt(9)
                run_py.font.color.rgb = RGBColor(0, 128, 128)
                
            # Meaning
            run_mean = p.add_run(f" - {c.get('meaning', '')}")
            run_mean.font.name = base_font
            run_mean.font.size = Pt(9.5)
            run_mean.font.color.rgb = RGBColor(64, 64, 64)

    doc.save(filename)


# ==========================================
# 🛠️ 무적의 API 재시도 (Robust Retry) 로직
# ==========================================
async def robust_generate_content(client, model, contents, config, max_retries=10):
    import re
    
    for attempt in range(max_retries):
        try:
            # 동기/비동기 혼용 시 발생할 수 있는 문제를 피하기 위해 비동기로 실행
            loop = asyncio.get_running_loop()
            response = await loop.run_in_executor(
                None,
                lambda: client.models.generate_content(
                    model=model,
                    contents=contents,
                    config=config
                )
            )
            return response
        except Exception as e:
            error_msg = str(e)
            if attempt == max_retries - 1:
                print(f"❌ 최대 재시도 횟수({max_retries}회)를 초과하여 최종 실패했습니다.")
                raise e
            
            # 기본 대기 시간: 시도 횟수에 따라 30초, 60초, 120초 ...
            wait_time = 30 * (2 ** attempt) 
            
            # 429 에러 메시지에 'retry in 33.23s'가 포함되어 있다면 해당 시간 추출
            retry_match = re.search(r'retry in ([\d\.]+)s', error_msg, re.IGNORECASE)
            if retry_match:
                suggested_wait = float(retry_match.group(1))
                wait_time = max(wait_time, suggested_wait + 5) # 제시된 시간보다 5초 넉넉하게 대기
            
            # 최대 대기 시간은 300초(5분)로 제한
            wait_time = min(wait_time, 300)
            
            print(f"⚠️ API 호출 실패 ({attempt+1}/{max_retries}): {error_msg.splitlines()[0]}")
            print(f"⏳ 트래픽 초과(Free Tier) 또는 서버 부하로 인해 {wait_time:.1f}초 대기 후 재시도합니다...")
            await asyncio.sleep(wait_time)


# ==========================================
# 🎵 메인 실행
# ==========================================
async def main_async():
    client = genai.Client(api_key=GEMINI_API_KEY)
    
    # 케냐 시간(EAT) 기준 날짜 계산 (UTC+3)
    import pytz
    local_tz = pytz.timezone("Africa/Nairobi")
    now = datetime.datetime.now(local_tz)
    weekdays = ["월", "화", "수", "목", "금", "토", "일"]
    weekday_str = weekdays[now.weekday()]
    today_folder_name = f"{now.strftime('%Y%m%d')} ({weekday_str})"
    
    # 출력 폴더 결정
    if IS_GITHUB_ACTIONS:
        # GitHub Actions: 임시 로컬 폴더에 생성 후 Google Drive에 업로드
        base_dir = os.path.join(os.getcwd(), "output")
    else:
        # 로컬 실행: 기존과 동일하게 스크립트 위치 기준
        base_dir = os.path.dirname(os.path.abspath(__file__))
    
    target_folder = os.path.join(base_dir, today_folder_name)
    os.makedirs(target_folder, exist_ok=True)

    print(f"\n🔄 [다국어 자동화 파이프라인 시작]")
    print(f"👉 오늘 생성 방식: 100% 무한 랜덤 시나리오 (One-Shot 일격필살 생성)")
    if IS_GITHUB_ACTIONS:
        print(f"☁️ GitHub Actions 클라우드 실행 모드")

    # 🎲 오늘의 시나리오를 SCENARIO_POOL에서 물리적으로 강제 추출
    today_scenario = random.choice(SCENARIO_POOL)
    scenario_category = today_scenario[0]
    scenario_detail = today_scenario[1]
    print(f"🎲 오늘의 강제 추출 시나리오: [{scenario_category}]")
    print(f"   📌 상황: {scenario_detail}")

    print("⏳ 단 한 번의 API 호출로 한국어, 영어, 중국어 학습 스크립트를 동시에 생성합니다...")
    master_prompt = f"""
    당신은 언어 학습자를 위한 최고의 콘텐츠 크리에이터입니다.
    오늘의 시나리오는 시스템이 물리적으로 무작위 추출하여 강제 지정한 주제입니다. 반드시 이 주제와 상황으로만 작성하세요.

    [🎯 오늘의 강제 지정 시나리오]
    - 대분류: {scenario_category}
    - 구체적 상황: {scenario_detail}

    [구어체 & 자연스러운 표현 최우선 원칙]
    이 스크립트의 최종 목표는 "한국식 암기 공부로는 절대 배울 수 없는 원어민의 찐 구어체 표현과 뉘앙스"를 체득하는 것입니다.
    - 교과서적이고 딱딱한 문어체(Written style)는 절대 쓰지 마세요.
    - 실제 원어민이 친구에게 말하듯 자연스러운 구어체(Colloquial/Spoken style)만 사용하세요.
    - 영어: 구동사(Phrasal Verbs), 숙어(Idioms), 추임새(like, I mean, honestly, literally 등)를 적극 반영하세요.
    - 중국어: 위챗 메신저나 일상에서 쓰는 찐 구어체(口语化)와 자연스러운 어기조사(嘛, 吧, 呀, 啊, 呢, 嗯 등)를 반영하세요.
    - 한국어 번역도 어색한 직역("나는 그것에 동의했다")이 아니라, 한국인이 실제로 입에서 나오는 리얼한 말투("그냥 쿨하게 오케이해줬어")로 작성하세요.

    [작성 규칙]
    - 기계적이거나 뻔한 상황(예: "요즘 바빠요", "잘 지내요?")은 절대 피하세요. 위에서 지정된 상황 안에서 매우 구체적이고 생생한 디테일을 설정하세요.
    - 한국어 원문은 상황에 완벽히 맞는 톤 앤 매너를 유지하세요.
    - 영어: 미국식 영어 기준. Chunking(끊어 읽기 표시인 '/' 또는 '//') 절대 사용 금지.
    - 중국어: 중국어 간체자로 번역. 모든 중국어 원문과 중국어 단어에 발음 기호인 병음(Pinyin) 필수 표기. Chunking 절대 사용 금지.
    - 분량: 가장 핵심이 되는 실전 문장 5개로 구성.

    [🚨초중요 - TTS 음성 합성용🚨]
    만약 대화문 형식이더라도, 각 언어별 문장 앞에는 'A:', 'B:', '이름:', '남:', '여:' 등 **어떠한 화자 표시나 기호도 절대 붙이지 마세요**. (TTS가 이름을 그대로 읽어버리는 치명적인 오류가 발생합니다.) 
    따옴표나 끊어 읽기 표시('/', '//')도 절대 쓰지 말고 순수한 대사 텍스트만 작성해야 합니다.

    응답은 반드시 아래 JSON 형식으로만 작성해주세요:
    {{
      "background_kr": "어떤 상황, 어떤 주제, 어떤 형식인지에 대한 구체적인 배경 설명 (한글로 2~3문장)",
      "english": {{
        "sentences": [
          {{
            "original": "영어 번역 원문 1",
            "translation": "한국어 원문 1"
          }}
        ],
        "vocabulary": [
          {{"word": "핵심단어", "pos": "품사", "meaning": "뜻"}}
        ],
        "collocations": [
          {{"expression": "덩어리 표현", "meaning": "뜻"}}
        ]
      }},
      "chinese": {{
        "sentences": [
          {{
            "original": "중국어 번역 원문 1",
            "pinyin": "병음",
            "translation": "한국어 원문 1"
          }}
        ],
        "vocabulary": [
          {{"word": "핵심단어", "pinyin": "병음", "pos": "품사", "meaning": "뜻"}}
        ],
        "collocations": [
          {{"expression": "덩어리 표현", "pinyin": "병음", "meaning": "뜻"}}
        ]
      }}
    }}
    """
    
    master_data = None
    try:
        response = await robust_generate_content(
            client=client,
            model='gemini-flash-latest',
            contents=master_prompt,
            config=types.GenerateContentConfig(response_mime_type="application/json")
        )
        resp_text = response.text.strip()
        if resp_text.startswith("```json"):
            resp_text = resp_text[7:]
        if resp_text.endswith("```"):
            resp_text = resp_text[:-3]
        master_data = json.loads(resp_text.strip())
        
        # 화자 접두사 강력 제거 (A:, 이름:, [A], A말하기: 등) 및 끊어읽기(/) 기호 완벽 제거
        import re
        for lang_key in ['english', 'chinese']:
            if lang_key in master_data and 'sentences' in master_data[lang_key]:
                for s in master_data[lang_key]['sentences']:
                    for key in ['original', 'translation']:
                        if key in s:
                            text = s[key]
                            text = re.sub(r'^[A-Za-z0-9가-힣一-龥\s\-\_]{1,15}[:：]\s*', '', text)
                            text = re.sub(r'^[\(\[][A-Za-z0-9가-힣一-龥\s\-\_]{1,15}[\)\]]\s*', '', text)
                            # AI가 지시를 무시하고 넣은 끊어읽기 기호(/) 강제 삭제
                            text = text.replace('/', '').replace('//', '')
                            s[key] = text.strip()

        # pypinyin을 이용한 병음 생성 및 교정
        from pypinyin import pinyin, Style
        if 'chinese' in master_data:
            cd = master_data['chinese']
            
            def fix_pinyin(text):
                if not text: return ""
                py_list = pinyin(text, style=Style.TONE)
                return " ".join([item[0] for item in py_list])
            
            if 'sentences' in cd:
                for s in cd['sentences']:
                    if 'original' in s:
                        s['pinyin'] = fix_pinyin(s['original'])
            if 'vocabulary' in cd:
                for v in cd['vocabulary']:
                    if 'word' in v:
                        v['pinyin'] = fix_pinyin(v['word'])
            if 'collocations' in cd:
                for c in cd['collocations']:
                    if 'expression' in c:
                        c['pinyin'] = fix_pinyin(c['expression'])

    except Exception as e:
        print(f"❌ 다국어 스크립트 One-Shot 생성에 최종 실패하여 작업을 종료합니다: {e}")
        return

    print("✅ 다국어 스크립트 One-Shot 생성 완료!")

    for lang in LANGUAGES:
        print(f"\n=========================================")
        print(f"▶️ [{lang['name']}] 파일 생성 시작...")
        
        # 언어 식별자
        lang_key = 'english' if lang['name'] == '영어' else 'chinese'
        data = master_data.get(lang_key)
        
        if not data:
            print(f"❌ {lang['name']} 데이터가 JSON 응답에 없습니다. 다음 언어로 넘어갑니다.")
            continue
            
        # background_kr은 공통 항목이므로 복사
        data['background_kr'] = master_data.get('background_kr', '')

        # 워드 파일 저장
        doc_filename = os.path.join(target_folder, f"{lang['name']}_학습자료.docx")
        try:
            create_docx(data, doc_filename, lang['name'])
            print(f"✅ 문서 파일 저장 완료: {doc_filename}")
        except Exception as e:
            print(f"❌ 문서 파일 저장 중 오류 발생: {e}")

        # 음성 파일 저장 (대화 형식이므로 남녀 성우 교대로 생성하여 바이너리 병합)
        print(f"🎵 {lang['name']} 음성(TTS) 파일 생성 중 (남녀 듀엣 버전)...")
        audio_filename = os.path.join(target_folder, f"{lang['name']}_학습음성.mp3")
        try:
            temp_files = []
            sentences = [s.get('original', '') for s in data.get('sentences', [])]
            
            for i, sent_text in enumerate(sentences):
                selected_voice = lang['voices'][i % 2]
                temp_fn = f"{audio_filename}_temp_{i}.mp3"
                
                communicate = edge_tts.Communicate(sent_text, selected_voice, rate=lang['rate'])
                await communicate.save(temp_fn)
                temp_files.append(temp_fn)
            
            with open(audio_filename, 'wb') as outfile:
                for temp_fn in temp_files:
                    with open(temp_fn, 'rb') as infile:
                        outfile.write(infile.read())
                    try:
                        os.remove(temp_fn)
                    except:
                        pass
                        
            print(f"✅ 듀엣 음성 파일 저장 완료: {audio_filename}")
        except Exception as e:
            print(f"❌ 음성 파일 저장 중 오류 발생: {e}")

    # GitHub Actions 환경이면 Google Drive에 업로드
    if IS_GITHUB_ACTIONS:
        print("\n☁️ Google Drive에 파일 업로드 중...")
        try:
            service = get_drive_service()
            if service:
                upload_folder_to_drive(service, target_folder, today_folder_name)
                print("✅ Google Drive 업로드 완료!")
            else:
                print("⚠️ Drive 서비스를 초기화할 수 없어 업로드를 건너뜁니다.")
        except Exception as e:
            print(f"❌ Google Drive 업로드 실패: {e}")

    print("\n🎉 모든 언어(영어/중국어)의 작업이 완벽하게 완료되었습니다!")

if __name__ == "__main__":
    asyncio.run(main_async())
